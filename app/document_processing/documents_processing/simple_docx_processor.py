#!/usr/bin/env python3
import sys
from pathlib import Path
from docx import Document
import json

def process_docx_simple(file_path, output_file):
    """Simple DOCX processor using python-docx directly"""
    try:
        print(f"📝 Processing DOCX: {file_path}")
        
        # Load document
        doc = Document(str(file_path))
        
        # Extract text
        content_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect headings
                style_name = para.style.name.lower()
                if 'heading' in style_name:
                    level = 1
                    try:
                        level = int(style_name.split()[-1])
                    except:
                        pass
                    content_parts.append(f"\n{'#' * level} {text}\n")
                else:
                    content_parts.append(text)
        
        # Extract tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            content_parts.append(f"\n\n## Table {table_count}\n")
            
            # Convert table to markdown
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            
            if rows:
                # Header
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                content_parts.append(header)
                content_parts.append(separator)
                
                # Data rows
                for row in rows[1:]:
                    data_row = "| " + " | ".join(row) + " |"
                    content_parts.append(data_row)
            
            content_parts.append("\n")
        
        # Combine content
        final_content = "\n".join(content_parts)
        
        # Save to markdown
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(final_content)
        
        print(f"✅ Success! Content: {len(final_content)} chars, Tables: {table_count}")
        print(f"📄 Saved to: {output_file}")
        
        return {
            "success": True,
            "content_length": len(final_content),
            "tables_count": table_count,
            "output_file": str(output_file)
        }
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python simple_docx_processor.py <input_file> <output_file>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    result = process_docx_simple(input_file, output_file)
    
    # Print result as JSON for parsing
    print("RESULT_START")
    print(json.dumps(result))
    print("RESULT_END")
